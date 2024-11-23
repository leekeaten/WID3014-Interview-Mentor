import streamlit as st
from jamaibase import JamAI, protocol as p
import os
from docx import Document
from io import BytesIO
import random
import string
from PyPDF2 import PdfReader
from dotenv import load_dotenv

# Load environment variables from the .env file
load_dotenv()

api_key = os.getenv('JAMAI_API_KEY')
project_id = os.getenv('JAMAI_PROJECT_ID')

jamai = JamAI(api_key=api_key, project_id=project_id)


# Function to extract text from PDF
def extract_text_from_pdf(pdf_file):
    pdf = PdfReader(pdf_file)
    text = ""
    for page in pdf.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text


# Function to generate a random filename
def generate_random_filename(extension=".docx"):
    random_str = ''.join(random.choices(string.ascii_letters + string.digits, k=10))
    return f"interview_questions_{random_str}{extension}"


# Set up the Streamlit app
st.set_page_config(page_title="Interview Mentor", page_icon="📝")
st.title("🌟 Interview Mentor - Your Partner for Interview Excellence")

# Custom CSS to style the UI
st.markdown(
    """
    <style>
    body {
        background-color: #1e1e1e;
        color: #f0f0f0;
    }
    .generated-output {
        background-color: #444;
        padding: 15px;
        border-radius: 10px;
        margin-top: 20px;
        box-shadow: 0px 4px 10px rgba(0, 0, 0, 0.5);
        color: #f0f0f0;
    }
    .generated-output h4 {
        color: #FFA500;
    }
    .stButton>button {
        background-color: #4CAF50;
        color: white;
        border: none;
        border-radius: 5px;
        padding: 10px 20px;
        cursor: pointer;
    }
    .stButton>button:hover {
        background-color: #45a049;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Containers for inputs
with st.container():
    st.header("📄 Upload CV and Enter Target Job Role")
    # Upload PDF CV
    cv_pdf = st.file_uploader("Upload CV (PDF format)", type="pdf")
    # Job Role input
    job_role = st.text_area("✍️ Enter Job Role")

# Action to process inputs
if st.button("🚀 Generate Interview Questions", use_container_width=True):
    if cv_pdf and job_role:
        # Extract text from CV PDF
        cv_text = extract_text_from_pdf(cv_pdf)

        # Add rows to the existing table with the input data
        try:
            completion = jamai.add_table_rows(
                "action",
                p.RowAddRequest(
                    table_id="interview-helper",
                    data=[{"cv": cv_text, "job_role": job_role}],
                    stream=False
                )
            )

            # Display the output generated in the columns
            if completion.rows:
                output_row = completion.rows[0].columns
                interview_questions = output_row.get("interview_questions")
                answer_suggestions = output_row.get("answer_suggestions")
                summary = output_row.get("summary")

                st.subheader("✨ Generated Interview Questions")
                st.markdown(
                    f"""
                    <div class="generated-output">
                        <h4>❓ Interview Answers:</h4> <p>{interview_questions.text if interview_questions else 'N/A'}</p>
                        <h4>💡 Suggested Answers:</h4> <p>{answer_suggestions.text if answer_suggestions else 'N/A'}</p>
                    </div>
                    """,
                    unsafe_allow_html=True
                )

                # Download the interview preparation report as a .docx file
                with st.container():
                    st.subheader("📥 Download Interview Preparation Document")
                    doc = Document()
                    doc.add_heading("Interview Preparation Report", level=1)
                    # Self Introduction Section
                    doc.add_heading("Self-Introduction", level=2)
                    doc.add_paragraph(summary.text if summary else 'N/A')
                    # Questions Section
                    doc.add_heading("Interview Questions", level=2)
                    doc.add_paragraph(interview_questions.text if interview_questions else 'N/A')
                    # Answers
                    doc.add_heading("Suggested Answers", level=2)
                    doc.add_paragraph(answer_suggestions.text if answer_suggestions else 'N/A')
                    

                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    st.download_button(
                        label="📄 Download Document as .docx",
                        data=buffer,
                        file_name=generate_random_filename(),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.error("⚠️ Failed to get a response. Please try again.")
        except Exception as e:
            st.error(f"❌ An error occurred: {e}")
    else:
        st.warning("⚠️ Please upload a CV and enter a job role.")